"""Rent-review pack generator.

Given a Lease + LeaseEvent + comparables, calls Claude to produce a structured
pack (headline numbers + four documents in markdown), then renders each
document to a Word .docx for the surveyor to edit and send.
"""

from __future__ import annotations

import json
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from anthropic import Anthropic
from pydantic import BaseModel, Field, ValidationError

from .extract import DEFAULT_MODEL
from .prompts import PACK_SYSTEM_PROMPT


PACK_TOOL_NAME = "record_pack"


# ---- structured output schema ------------------------------------------


class PackOutput(BaseModel):
    """The structured payload Claude returns via the `record_pack` tool."""

    recommended_opening_gbp: float = Field(description="Opening position to send to the tenant, in GBP per annum.")
    recommended_settlement_low_gbp: float = Field(description="Floor of the expected settlement range, GBP p.a.")
    recommended_settlement_high_gbp: float = Field(description="Ceiling of the expected settlement range, GBP p.a.")

    landlord_memo_markdown: str
    comparables_schedule_markdown: str
    itza_analysis_markdown: str
    trigger_letter_markdown: str

    confidence_notes: str | None = Field(
        default=None,
        description="Any caveats or confidence notes the surveyor should be aware of (low evidence count, unusual mechanism etc.).",
    )


@dataclass
class GeneratedPack:
    output: PackOutput
    model: str
    elapsed_seconds: float
    input_tokens: int
    output_tokens: int
    cache_read_tokens: int
    cache_write_tokens: int


# ---- main entry point ---------------------------------------------------


def generate_pack(
    *,
    lease_record: dict[str, Any],
    event_summary: dict[str, Any],
    comparables: list[dict[str, Any]],
    model: str | None = None,
) -> GeneratedPack:
    """Call Claude to produce the pack content. Pure — does no IO except the API call."""
    client = Anthropic()
    chosen_model = model or DEFAULT_MODEL

    user_text = _format_user_message(lease_record, event_summary, comparables)
    tool = _build_tool()

    started = time.monotonic()
    response = client.messages.create(
        model=chosen_model,
        max_tokens=8000,
        system=[
            {
                "type": "text",
                "text": PACK_SYSTEM_PROMPT,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        tools=[tool],
        tool_choice={"type": "tool", "name": PACK_TOOL_NAME},
        messages=[{"role": "user", "content": [{"type": "text", "text": user_text}]}],
    )
    elapsed = time.monotonic() - started

    tool_input = _extract_tool_input(response)
    try:
        output = PackOutput.model_validate(tool_input)
    except ValidationError as exc:
        raise RuntimeError(
            f"Pack generator returned malformed JSON: {exc}\n\n"
            f"Raw input:\n{json.dumps(tool_input, indent=2, default=str)[:2000]}"
        ) from exc

    usage = response.usage
    return GeneratedPack(
        output=output,
        model=chosen_model,
        elapsed_seconds=elapsed,
        input_tokens=getattr(usage, "input_tokens", 0),
        output_tokens=getattr(usage, "output_tokens", 0),
        cache_read_tokens=getattr(usage, "cache_read_input_tokens", 0) or 0,
        cache_write_tokens=getattr(usage, "cache_creation_input_tokens", 0) or 0,
    )


# ---- prompt assembly ----------------------------------------------------


def _format_user_message(
    lease_record: dict[str, Any],
    event_summary: dict[str, Any],
    comparables: list[dict[str, Any]],
) -> str:
    parts: list[str] = []
    parts.append("# Subject lease — structured abstraction\n")
    parts.append("```json\n" + json.dumps(lease_record, indent=2, default=str) + "\n```\n")

    parts.append("\n# This rent-review event\n")
    parts.append("```json\n" + json.dumps(event_summary, indent=2, default=str) + "\n```\n")

    parts.append(f"\n# Comparable evidence ({len(comparables)} items)\n")
    if not comparables:
        parts.append(
            "*No comparable evidence has been added yet. Produce the memo and "
            "trigger letter on the basis of the passing rent and the lease's "
            "review mechanism, and explicitly flag in the comparables schedule "
            "and ITZA analysis that no market evidence was supplied.*\n"
        )
    else:
        parts.append("```json\n" + json.dumps(comparables, indent=2, default=str) + "\n```\n")

    parts.append(
        "\nNow call the `record_pack` tool with the headline numbers and the "
        "four documents. Use clean markdown that will render well in a Word "
        "document (headings, tables, bullet lists)."
    )
    return "\n".join(parts)


def _build_tool() -> dict:
    return {
        "name": PACK_TOOL_NAME,
        "description": "Record the complete rent-review pack content.",
        "input_schema": PackOutput.model_json_schema(),
        "cache_control": {"type": "ephemeral"},
    }


def _extract_tool_input(response) -> dict:
    for block in response.content:
        if getattr(block, "type", None) == "tool_use" and block.name == PACK_TOOL_NAME:
            return block.input
    raise RuntimeError(
        f"Model did not call the {PACK_TOOL_NAME!r} tool. Stop reason: {response.stop_reason}"
    )


# ---- Word .docx rendering -----------------------------------------------


def _firm_template_path(kind: str) -> Path | None:
    """If the firm has uploaded a Word template for `kind`, return its path.
    Otherwise return None (caller falls back to the LeaseOS default style).

    Templates live next to the documents store at `data/templates/{kind}.docx`
    so they're co-located with everything else file-system-y. Single-tenant
    convention for v1; will move to per-firm storage once Clerk Organisations
    lands.
    """
    from .api.config import get_settings

    settings = get_settings()
    templates_dir = settings.storage_dir.parent / "templates"
    candidate = templates_dir / f"{kind}.docx"
    return candidate if candidate.exists() else None


def render_docx(*, kind: str, markdown_content: str, output_path: Path, title: str) -> None:
    """Render a markdown document to .docx. Lightweight markdown subset:
    H1/H2/H3, bold inline (**), bullet lists, tables (pipe syntax), paragraphs.

    If the firm has uploaded a custom template at `data/templates/{kind}.docx`,
    that template's first-page header / styles / firm logo are preserved and
    the generated content is appended below them. Otherwise the LeaseOS
    default Calibri 11pt style is used.
    """
    from docx import Document
    from docx.shared import Pt, Cm

    template_path = _firm_template_path(kind)
    if template_path is not None:
        # Open the firm's template — keeps their styles, logos, footers.
        doc = Document(str(template_path))
        # Don't touch styles or margins — assume the firm set those deliberately.
    else:
        doc = Document()
        # LeaseOS default style
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        for section in doc.sections:
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)

    if title:
        h = doc.add_heading(title, level=0)
        if h.runs:
            h.runs[0].font.size = Pt(16)

    _render_markdown_into(doc, markdown_content)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _render_markdown_into(doc, md: str) -> None:
    """Tiny markdown → docx renderer. Not a full parser — handles what the
    pack-generator prompt actually emits."""
    from docx.shared import Pt

    lines = md.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Tables — markdown pipe syntax
        if line.lstrip().startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            table_lines = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _render_table(doc, table_lines)
            continue

        # Bullet lists
        if line.lstrip().startswith(("- ", "* ")):
            while i < len(lines) and lines[i].lstrip().startswith(("- ", "* ")):
                bullet_text = lines[i].lstrip()[2:]
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, bullet_text)
                i += 1
            continue

        # Numbered lists
        if line.lstrip()[:2].rstrip(".").isdigit() and ". " in line:
            while i < len(lines) and lines[i].lstrip()[:2].rstrip(".").isdigit() and ". " in lines[i]:
                num_text = lines[i].split(". ", 1)[1]
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, num_text)
                i += 1
            continue

        # Plain paragraph (collect until blank line)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].lstrip().startswith(("#", "-", "*", "|")):
            para_lines.append(lines[i].rstrip())
            i += 1
        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(para_lines))


def _render_table(doc, table_lines: list[str]) -> None:
    """Render a markdown pipe-table. Skips the separator line."""
    rows: list[list[str]] = []
    for li, raw in enumerate(table_lines):
        if li == 1:  # separator
            continue
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = ""  # clear default
            p = cell.paragraphs[0]
            _add_inline_runs(p, text)
            if ri == 0:
                for run in p.runs:
                    run.bold = True


def _add_inline_runs(paragraph, text: str) -> None:
    """Handle **bold** and plain text in a single paragraph. Keep it simple."""
    parts = text.split("**")
    for idx, segment in enumerate(parts):
        run = paragraph.add_run(segment)
        run.bold = (idx % 2 == 1)
