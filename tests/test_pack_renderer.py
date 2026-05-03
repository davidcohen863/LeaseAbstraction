"""Tests for the pack-generator's markdown → .docx renderer.

The Anthropic call in `pack_generator.generate_pack` is hard to test without
mocking; the renderer (`render_docx`, `_render_markdown_into`, `_render_table`,
`_add_inline_runs`) is pure local logic and very testable.

What we care about:
  * Output is a real .docx (zip with the right inner files) — `python-docx`
    can re-open it without complaining
  * Headings, bullet lists, paragraphs, tables, and **bold** all survive
    the round-trip
  * Empty markdown doesn't crash
  * Tables with ragged-row counts don't crash and produce the right column
    count
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as DocxDocument

from leaseos.pack_generator import (
    _add_inline_runs,
    _render_markdown_into,
    _render_table,
    render_docx,
)


def _read(path: Path):
    """Re-open the .docx and return (paragraphs, tables) for assertion."""
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    tables = [
        [[cell.text for cell in row.cells] for row in t.rows] for t in doc.tables
    ]
    return paragraphs, tables


class TestRenderDocx:
    def test_minimal_renders_and_reopens(self, tmp_path):
        out = tmp_path / "memo.docx"
        render_docx(
            kind="landlord_memo",
            markdown_content="Hello world.",
            output_path=out,
            title="Landlord memo",
        )
        assert out.exists()
        paragraphs, _ = _read(out)
        # Title appears as a heading paragraph + body paragraph.
        assert "Landlord memo" in paragraphs
        assert "Hello world." in paragraphs

    def test_headings_levels(self, tmp_path):
        out = tmp_path / "h.docx"
        md = "# H1\n\n## H2\n\n### H3\n\nbody text"
        render_docx(kind="landlord_memo", markdown_content=md, output_path=out, title="")
        paragraphs, _ = _read(out)
        # All three headings + the body paragraph survive.
        assert "H1" in paragraphs
        assert "H2" in paragraphs
        assert "H3" in paragraphs
        assert "body text" in paragraphs

    def test_bullet_list_each_item_separate_paragraph(self, tmp_path):
        out = tmp_path / "bl.docx"
        md = "- alpha\n- beta\n- gamma\n"
        render_docx(kind="landlord_memo", markdown_content=md, output_path=out, title="")
        paragraphs, _ = _read(out)
        # Each bullet becomes its own paragraph.
        assert "alpha" in paragraphs
        assert "beta" in paragraphs
        assert "gamma" in paragraphs

    def test_bold_inline_marker_stripped_in_body_text(self, tmp_path):
        out = tmp_path / "b.docx"
        md = "Opening rent of **£58,000** is recommended."
        render_docx(kind="landlord_memo", markdown_content=md, output_path=out, title="")
        paragraphs, _ = _read(out)
        # The literal ** marker should be stripped — only the text remains.
        body = next(p for p in paragraphs if "58,000" in p)
        assert "**" not in body
        assert "£58,000" in body

    def test_empty_markdown_does_not_crash(self, tmp_path):
        out = tmp_path / "empty.docx"
        render_docx(
            kind="landlord_memo", markdown_content="", output_path=out, title="Title only"
        )
        assert out.exists()
        paragraphs, _ = _read(out)
        assert "Title only" in paragraphs

    def test_creates_parent_directory(self, tmp_path):
        # If the caller passes data/packs/<uuid>/memo.docx and the dir
        # doesn't exist yet, the renderer should create it (mirrors the
        # pack_worker call site).
        out = tmp_path / "deep" / "nested" / "dir" / "memo.docx"
        render_docx(kind="landlord_memo", markdown_content="x", output_path=out, title="")
        assert out.exists()


class TestRenderTable:
    def test_pipe_table_column_count_and_header(self, tmp_path):
        md = (
            "## Comparables\n\n"
            "| Address | Rent | Area |\n"
            "|---|---|---|\n"
            "| 12 High St | £50,000 | 1200 |\n"
            "| 14 High St | £52,500 | 1250 |\n"
        )
        out = tmp_path / "t.docx"
        render_docx(kind="comparables_schedule", markdown_content=md, output_path=out, title="")
        _, tables = _read(out)
        assert len(tables) == 1
        rows = tables[0]
        # Header + 2 data rows; the markdown separator line is dropped.
        assert len(rows) == 3
        assert len(rows[0]) == 3
        assert rows[0] == ["Address", "Rent", "Area"]
        assert rows[1][1] == "£50,000"
        assert rows[2][2] == "1250"

    def test_ragged_row_does_not_crash(self, tmp_path):
        # A model could plausibly emit a row with a missing trailing cell.
        # The renderer should pad rather than blow up.
        md = (
            "| A | B | C |\n"
            "|---|---|---|\n"
            "| 1 | 2 |\n"
            "| x | y | z |\n"
        )
        out = tmp_path / "r.docx"
        render_docx(kind="comparables_schedule", markdown_content=md, output_path=out, title="")
        _, tables = _read(out)
        assert len(tables) == 1
        # All rows padded to 3 columns.
        assert all(len(row) == 3 for row in tables[0])
        # Missing cell shows as empty string (not None / not a crash).
        assert tables[0][1][2] == ""

    def test_table_lines_only(self, tmp_path):
        # Direct test of _render_table — no surrounding markdown.
        from docx import Document

        doc = Document()
        _render_table(
            doc,
            [
                "| col1 | col2 |",
                "|---|---|",
                "| a | b |",
            ],
        )
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "col1"
        assert doc.tables[0].rows[1].cells[1].text == "b"

    def test_empty_table_lines_no_crash(self):
        from docx import Document

        doc = Document()
        # Just the separator — no header, no rows. Should silently no-op.
        _render_table(doc, [])
        assert len(doc.tables) == 0


class TestInlineRuns:
    def test_alternating_bold_segments(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph()
        _add_inline_runs(p, "plain **bold** plain **more bold** end")
        runs = p.runs
        # 5 segments split on '**': plain / bold / plain / more bold / end
        assert len(runs) == 5
        assert runs[0].bold is False
        assert runs[1].bold is True
        assert runs[1].text == "bold"
        assert runs[3].bold is True
        assert runs[3].text == "more bold"
        assert runs[4].bold is False

    def test_no_bold_marker(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph()
        _add_inline_runs(p, "no markers here")
        assert len(p.runs) == 1
        assert p.runs[0].bold is False
        assert p.runs[0].text == "no markers here"
