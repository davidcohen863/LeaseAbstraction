"""Tests for the per-firm Word template endpoints + the pack_generator hook
that picks them up.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from leaseos.api.routes.templates import KIND_LABELS, VALID_KINDS, _template_path, _templates_dir
from leaseos.pack_generator import _firm_template_path, render_docx


def _make_docx_bytes(title: str = "Firm letterhead") -> bytes:
    """Build a minimal valid .docx in-memory for upload tests."""
    doc = DocxDocument()
    doc.add_paragraph(title)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestValidKinds:
    def test_kinds_match_pack_documents(self):
        assert VALID_KINDS == {
            "landlord_memo",
            "comparables_schedule",
            "itza_analysis",
            "trigger_letter",
        }

    def test_kind_labels_cover_every_kind(self):
        for k in VALID_KINDS:
            assert k in KIND_LABELS


class TestRouteContract:
    def test_unknown_kind_rejected(self, test_client):
        resp = test_client.post("/templates/not_a_real_kind", files={"file": ("x.docx", b"x")})
        assert resp.status_code == 400

    def test_non_docx_rejected(self, test_client):
        resp = test_client.post(
            "/templates/landlord_memo",
            files={"file": ("notes.pdf", b"%PDF-1.4")},
        )
        assert resp.status_code == 400

    def test_garbage_docx_rejected(self, test_client):
        # Right extension, wrong contents — python-docx will fail to parse.
        resp = test_client.post(
            "/templates/landlord_memo",
            files={"file": ("evil.docx", b"this is not a word document")},
        )
        assert resp.status_code == 400
        assert "parse" in resp.text.lower() or "word" in resp.text.lower()

    def test_upload_then_list_then_delete_round_trip(self, test_client, tmp_path, monkeypatch):
        # Point storage_dir at tmp so we don't touch the dev DB's data/
        from leaseos.api.config import get_settings

        settings = get_settings()
        monkeypatch.setattr(settings, "storage_dir", tmp_path / "documents")

        # 1. baseline: nothing uploaded
        r = test_client.get("/templates")
        assert r.status_code == 200
        assert all(t["uploaded"] is False for t in r.json())

        # 2. upload a real .docx
        payload = _make_docx_bytes(title="Acme Surveyors letterhead")
        r = test_client.post(
            "/templates/landlord_memo",
            files={"file": ("acme-style-memo.docx", payload)},
        )
        assert r.status_code == 201, r.text
        body = r.json()
        assert body["kind"] == "landlord_memo"
        assert body["uploaded"] is True
        assert body["original_filename"] == "acme-style-memo.docx"
        assert body["size_bytes"] == len(payload)

        # 3. list now reflects the upload
        r = test_client.get("/templates")
        landlord_memo = next(t for t in r.json() if t["kind"] == "landlord_memo")
        assert landlord_memo["uploaded"] is True
        # Other kinds still not uploaded
        others = [t for t in r.json() if t["kind"] != "landlord_memo"]
        assert all(t["uploaded"] is False for t in others)

        # 4. delete
        r = test_client.delete("/templates/landlord_memo")
        assert r.status_code == 204

        # 5. confirm gone
        r = test_client.get("/templates")
        assert all(t["uploaded"] is False for t in r.json())


class TestRenderDocxUsesTemplate:
    def test_falls_back_to_default_when_no_template(self, tmp_path, monkeypatch):
        from leaseos.api.config import get_settings

        settings = get_settings()
        monkeypatch.setattr(settings, "storage_dir", tmp_path / "documents")
        out = tmp_path / "out.docx"
        render_docx(
            kind="landlord_memo",
            markdown_content="Hello world",
            output_path=out,
            title="Memo",
        )
        assert out.exists()

    def test_uses_uploaded_template_as_base(self, tmp_path, monkeypatch):
        from leaseos.api.config import get_settings

        settings = get_settings()
        monkeypatch.setattr(settings, "storage_dir", tmp_path / "documents")

        # Drop a template that contains a unique marker string
        templates_dir = _templates_dir()
        template = templates_dir / "landlord_memo.docx"
        marker = "ACME-LETTERHEAD-MARKER-9f2a"
        DocxDocument().add_paragraph(marker)  # noqa
        from docx import Document as DocxDocument2

        doc = DocxDocument2()
        doc.add_paragraph(marker)
        doc.save(str(template))

        # Sanity: render finds the template
        assert _firm_template_path("landlord_memo") == template

        # Render new content — output should contain BOTH the marker and the new content
        out = tmp_path / "out.docx"
        render_docx(
            kind="landlord_memo",
            markdown_content="Recommended opening £67,000",
            output_path=out,
            title="Memo for Mr Patel",
        )
        rendered = DocxDocument2(str(out))
        text = "\n".join(p.text for p in rendered.paragraphs)
        assert marker in text, "Firm template's letterhead must survive into the output"
        assert "67,000" in text or "Recommended opening" in text, (
            "Generated content must be appended to the template"
        )

    def test_template_only_used_for_matching_kind(self, tmp_path, monkeypatch):
        # Uploading a landlord_memo template must NOT change trigger_letter renders.
        from leaseos.api.config import get_settings
        from docx import Document as DocxDocument2

        settings = get_settings()
        monkeypatch.setattr(settings, "storage_dir", tmp_path / "documents")
        templates_dir = _templates_dir()
        marker = "MEMO-ONLY-MARKER"
        doc = DocxDocument2()
        doc.add_paragraph(marker)
        doc.save(str(templates_dir / "landlord_memo.docx"))

        out = tmp_path / "letter.docx"
        render_docx(
            kind="trigger_letter",
            markdown_content="Body of trigger letter",
            output_path=out,
            title="Trigger letter",
        )
        rendered = DocxDocument2(str(out))
        text = "\n".join(p.text for p in rendered.paragraphs)
        assert marker not in text
